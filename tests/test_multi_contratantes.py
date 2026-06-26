"""Teste com múltiplas contratantes."""

from app import create_app
from app.services.context_builder import build_contract_context
from app.services.docx_generator import generate_docx
from docx import Document


def main():
    app = create_app()
    with app.app_context():
        form_data = {
            "tipo_servico": "assessoria_contabil",
            "inicio_vigencia": "2026-07-01",
            "forma_pagamento": "boleto",
            "dia_vencimento": "25",
            "mes_vencimento": "dia_fixo",
            "reajuste": "sm_ipca_igpm",
            "decima_terceira_parcela": "sim",
            "rescisao": "60_dias",
            "suspensao": "inadimplencia_2",
            "foro": "recife_pe",
            "detalhamento_servico": "escopo_padrao",
            "contratantes": [
                {
                    "cnpj": "42957943000157",
                    "enquadramento_tributario": "simples",
                    "razao_social": "Onvilla Ltda",
                    "valor_mensal": "5673.50",
                    "nome": "Andreia",
                    "cpf": "02175111989",
                },
                {
                    "cnpj": "21669413000133",
                    "enquadramento_tributario": "lucro_presumido",
                    "razao_social": "Ecovilla Ltda",
                    "valor_mensal": "5673.50",
                    "nome": "Hugo",
                    "cpf": "38232111968",
                },
            ],
            "empresa_contratada_id": 1,
            "cnpj_contratada": "24863877000174",
            "endereco_contratado": "Rua Irmã Maria David, 99",
            "nome_contratado": "André Abdom Ximenes Cavalcanti",
            "cpf_contratado": "05359403438",
            "crc_contratado": "PE-030717/O",
            "local": "Recife/PE",
            "data_contrato": "2026-06-15",
        }
        context = build_contract_context(form_data)
        path = generate_docx(context, "Grupo")
        doc = Document(str(path))
        table = doc.tables[1] if len(doc.tables) > 1 else doc.tables[0]
        data_rows = len(table.rows) - 1
        assert data_rows == 2, f"Esperado 2 linhas, obteve {data_rows}"
        print("Multi-contratante OK:", path)


if __name__ == "__main__":
    main()
